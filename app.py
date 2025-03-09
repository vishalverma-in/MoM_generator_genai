from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from llm import create_gemini_llm
from langchain_core.prompts import ChatPromptTemplate
import os

llm_gemini = create_gemini_llm()

def read_transcript_from_docx(file_path):
    """
    Read the transcript from a .docx file.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    doc = Document(file_path)
    transcript = ""
    for paragraph in doc.paragraphs:
        transcript += paragraph.text + "\n"
    return transcript.strip()

def preprocess_transcript(transcript):
    """
    Preprocess the transcript (if needed).
    """
    # Example: Remove extra spaces, filler words, etc.
    transcript = transcript.replace("\n", " ").strip()
    return transcript

def generate_mom(transcript):
    """
    Generate Minutes of the Meeting (MoM) using Gemini LLM.
    """
    prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            """
            You are an expert in summarizing meeting transcripts. Analyze the following transcript and generate Minutes of the Meeting (MoM). Include:
            1. Key discussion points.
            2. Decisions made.
            3. Action items with assigned owners.
            Format the output clearly with headings and bullet points. Include references to the speakers where relevant.
            """,
        ),
        ("human", "{transcript}"),
    ]
    )
    
    chain = prompt | llm_gemini
    
    response = chain.invoke(
        {
            "transcript": transcript,
        }
    )
    return response.content

def save_mom_to_file(mom_text, output_file="mom_output.docx"):
    """
    Save the MoM to a .docx file with proper formatting.
    """
    doc = Document()

    # Add title
    title = doc.add_heading("Minutes of the Meeting", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add MoM content
    for line in mom_text.split("\n"):
        if line.strip().startswith("#"):
            # Add heading
            doc.add_heading(line.strip("# ").strip(), level=2)
        else:
            # Add paragraph
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)

    # Save the document
    doc.save(output_file)
    print(f"MoM saved to {output_file}")

def main():
    # Path to the transcript .docx file
    transcript_file = r"V:\coforge_projects\genAI\MoM_generator\transcript.docx"

    # Read transcript from .docx file
    transcript = read_transcript_from_docx(transcript_file)

    # Preprocess transcript
    transcript = preprocess_transcript(transcript)

    # Generate MoM
    mom_text = generate_mom(transcript)
    print("Generated MoM:\n", mom_text)

    # Save MoM to file
    save_mom_to_file(mom_text)

if __name__ == "__main__":
    main()